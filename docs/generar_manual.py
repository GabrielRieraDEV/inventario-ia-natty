"""Genera el Manual de Usuario del sistema (Word .docx editable) con capturas.

Salida: docs/Manual_de_Usuario.docx
Uso:    python generar_manual.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Carpeta con las capturas (mockups de Google Stitch)
IMG = r"C:\Users\Gabriel Riera\Downloads\stitch_pa_donde_natty_smart_inventory (1)\stitch_pa_donde_natty_smart_inventory"

NAVY = RGBColor(0x12, 0x3A, 0x6B)
GOLD = RGBColor(0xC0, 0x9A, 0x4A)
GRAY = RGBColor(0x53, 0x60, 0x6E)

doc = Document()

# --- Estilo base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = s.bottom_margin = Inches(1.0)

_fig = [0]


def _color_heading(p, color=NAVY):
    for r in p.runs:
        r.font.color.rgb = color


def h1(text):
    p = doc.add_heading(text, level=1)
    _color_heading(p)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    _color_heading(p)
    return p


def para(text, justify=True, italic=False):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if italic:
        p.runs[0].font.italic = True
    return p


def pasos(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def vinetas(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def nota(text):
    p = doc.add_paragraph()
    r = p.add_run("Nota: ")
    r.font.bold = True
    r.font.color.rgb = GOLD
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.color.rgb = GRAY
    return p


def figura(carpeta, caption, portrait=False):
    _fig[0] += 1
    ruta = os.path.join(IMG, carpeta, "screen.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ruta):
        run = p.add_run()
        run.add_picture(ruta, width=Inches(2.7 if portrait else 6.2))
    else:
        r = p.add_run(f"[ Inserte aquí la captura: {carpeta} ]")
        r.font.italic = True
        r.font.color.rgb = GRAY
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f"Figura {_fig[0]}. {caption}")
    rc.font.size = Pt(9)
    rc.font.italic = True
    rc.font.color.rgb = GRAY
    doc.add_paragraph()


def toc():
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Haga clic derecho aquí y elija «Actualizar campos» para generar el índice."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, t, f3):
        run._r.append(e)


# =====================================================================
# PORTADA
# =====================================================================
for txt, sz, bold, col in [
    ("Universidad Nororiental Privada Gran Mariscal de Ayacucho", 13, True, NAVY),
    ("Escuela de Ingeniería en Informática — Núcleo Bolívar", 11, False, GRAY),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(sz); r.font.bold = bold; r.font.color.rgb = col

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MANUAL DE USUARIO"); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Digitalización y Registro de Productos con Inteligencia Artificial")
r.font.size = Pt(15); r.font.color.rgb = GRAY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bodegón Pa' Donde Natty C.A."); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = GOLD

for _ in range(6):
    doc.add_paragraph()

for txt, bold in [
    ("Autor: Gabriel Riera — C.I.: V-31.153.478", True),
    ("Tutor Académico: Ing. Carlos Moreno", False),
    ("Versión 1.0 — Ciudad Bolívar, 2026", False),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(11); r.font.bold = bold
doc.add_page_break()

# =====================================================================
# CONTENIDO
# =====================================================================
h1("Contenido")
toc()
doc.add_page_break()

# =====================================================================
# 1. INTRODUCCIÓN
# =====================================================================
h1("1. Introducción")
para("Este manual explica, paso a paso, cómo usar el Sistema de Digitalización y "
     "Registro de Productos con Inteligencia Artificial del Bodegón Pa' Donde Natty C.A. "
     "El sistema permite controlar el inventario de forma digital: registra entradas y "
     "salidas de mercancía, mantiene el stock actualizado en tiempo real, genera alertas "
     "y reportes, y —su función distintiva— reconoce los productos a partir de una "
     "fotografía tomada con el teléfono, agilizando el registro.")
para("El manual está dirigido a todo el personal que opera el sistema: administradores, "
     "gerentes y operarios de almacén. No requiere conocimientos técnicos avanzados.")

# =====================================================================
# 2. REQUISITOS
# =====================================================================
h1("2. Requisitos para usar el sistema")
vinetas([
    "Una computadora (PC o laptop) con conexión a Internet.",
    "Un navegador web actualizado (Google Chrome, Microsoft Edge o Firefox).",
    "Un teléfono inteligente con cámara y conexión a Internet, para la captura de fotos.",
    "Las credenciales de acceso (correo y contraseña) entregadas por el administrador.",
])
nota("La captura por foto desde el teléfono requiere que el sitio se abra mediante una "
     "dirección segura (https). El sistema desplegado ya cumple este requisito.")

# =====================================================================
# 3. ACCESO AL SISTEMA (LOGIN)
# =====================================================================
h1("3. Acceso al sistema")
para("Para ingresar al sistema:")
pasos([
    "Abra el navegador y escriba la dirección del sistema proporcionada por el administrador.",
    "En la pantalla de inicio de sesión, escriba su correo electrónico y su contraseña.",
    "Presione el botón «Iniciar sesión».",
    "Si los datos son correctos, accederá al Panel principal (Dashboard).",
])
figura("inicio_de_sesi_n_admin", "Pantalla de inicio de sesión.")
nota("Si la contraseña es incorrecta, el sistema mostrará «Credenciales inválidas». "
     "Verifique los datos o solicite ayuda al administrador. Se recomienda cambiar la "
     "contraseña inicial tras el primer ingreso.")

# =====================================================================
# 4. LA PANTALLA PRINCIPAL
# =====================================================================
h1("4. La pantalla principal y la navegación")
para("Tras iniciar sesión, la pantalla se organiza en tres zonas:")
vinetas([
    "Menú lateral (izquierda): acceso a todos los módulos (Panel general, Captura IA, "
    "Catálogo, Inventario, Movimientos, Alertas, Reportes y Configuración).",
    "Barra superior: buscador de inventario, campana de notificaciones (alertas activas), "
    "ícono de captura por QR y menú de su cuenta (configuración y cerrar sesión).",
    "Área central: el contenido del módulo seleccionado.",
])
figura("panel_principal_dashboard", "Panel principal (Dashboard) con indicadores y alertas.")
para("El Panel general (Dashboard) resume el estado del inventario: total de productos, "
     "alertas activas, movimientos del día y un vistazo a las existencias por categoría.")

# =====================================================================
# 5. CAPTURA POR IA (FLUJO QR)  -- la función estrella
# =====================================================================
h1("5. Registrar un producto con la cámara (Captura IA)")
para("Esta es la función principal del sistema: tomar una foto del producto con el "
     "teléfono para que la Inteligencia Artificial lo identifique y registre el movimiento "
     "(entrada o salida) automáticamente.")

h2("5.1. Vincular el teléfono mediante el código QR")
pasos([
    "En la computadora, abra «Captura IA» en el menú lateral (o el ícono de QR en la barra superior).",
    "Se mostrará un código QR en la pantalla.",
    "Con la cámara de su teléfono, escanee ese código QR.",
    "El teléfono abrirá automáticamente la pantalla de captura; la computadora quedará a la espera.",
])
figura("v_nculo_de_captura_qr", "Vinculación del teléfono mediante código QR.")
nota("El código QR es de un solo uso y caduca después de un tiempo. Si expira, recargue "
     "la pantalla en la computadora para generar uno nuevo.")

h2("5.2. Tomar la foto desde el teléfono")
pasos([
    "En el teléfono, permita el acceso a la cámara cuando se le solicite.",
    "Apunte la cámara al producto (o a su código de barras) dentro del recuadro guía.",
    "Presione el botón de captura. También puede elegir una foto de la galería.",
    "Espere unos segundos: la foto se envía y la IA reconoce el producto.",
])
figura("vista_m_vil_de_captura", "Vista de captura en el teléfono.", portrait=True)

h2("5.3. Confirmar el registro en la computadora")
para("Cuando la IA reconoce el producto, la computadora muestra el resultado para que "
     "usted lo revise y confirme:")
pasos([
    "Verifique el nombre del producto identificado; si es necesario, corríjalo.",
    "Revise la categoría y el nivel de confianza del reconocimiento.",
    "Indique la cantidad y, si corresponde, si es una entrada o una salida.",
    "Presione «Confirmar» para registrar el movimiento. El inventario se actualiza al instante.",
])
figura("resultado_ia_reconocimiento", "Resultado del reconocimiento por IA, listo para confirmar.")
nota("Para una salida, el producto debe existir ya en el inventario y haber existencias "
     "suficientes; de lo contrario el sistema lo advertirá.")

# =====================================================================
# 6. CATÁLOGO
# =====================================================================
h1("6. Catálogo de productos")
para("El Catálogo permite administrar los productos: crearlos, editarlos y darlos de baja.")
h2("6.1. Crear un producto")
pasos([
    "Abra «Catálogo» en el menú lateral.",
    "Presione «Nuevo producto».",
    "Complete los datos: nombre, código de barras, categoría, marca, presentación, "
    "precio, stock mínimo y, si aplica, la fecha de vencimiento.",
    "Presione «Guardar».",
])
h2("6.2. Editar o dar de baja un producto")
vinetas([
    "Para editar: pase el cursor sobre la fila del producto y presione el ícono de lápiz; "
    "modifique los datos y guarde.",
    "Para dar de baja: presione el ícono de papelera. El producto se desactiva, pero su "
    "historial de movimientos se conserva.",
    "También puede registrar un producto con la cámara desde el botón «Escanear con cámara».",
])
figura("cat_logo_de_productos", "Catálogo de productos.")

# =====================================================================
# 7. INVENTARIO
# =====================================================================
h1("7. Consultar el inventario")
para("El módulo Inventario muestra las existencias actuales de cada producto, calculadas "
     "automáticamente a partir de los movimientos registrados.")
pasos([
    "Abra «Inventario» en el menú lateral.",
    "Use la caja de búsqueda para filtrar por nombre o categoría.",
    "Observe el estado de cada producto: Normal, Bajo o Crítico, según su stock mínimo.",
])
figura("inventario_de_bodeg_n", "Inventario actual del bodegón.")
nota("También puede buscar desde la barra superior: al escribir y presionar Enter, el "
     "sistema lo lleva al inventario filtrado.")

# =====================================================================
# 8. MOVIMIENTOS
# =====================================================================
h1("8. Registrar entradas y salidas (manual)")
para("Además de la captura por foto, puede registrar movimientos manualmente:")
pasos([
    "Abra «Movimientos» en el menú lateral.",
    "Busque y seleccione el producto en el campo de búsqueda.",
    "Elija el tipo de movimiento: Entrada o Salida.",
    "Indique la cantidad y, opcionalmente, una nota.",
    "Revise el panel «Impacto del movimiento», que muestra el nuevo stock resultante.",
    "Presione «Confirmar movimiento».",
])
figura("registrar_movimiento", "Registro manual de un movimiento.")
nota("El sistema no permite que una salida deje el stock en negativo. También puede usar "
     "el botón de cámara para registrar la entrada o salida por foto.")

# =====================================================================
# 9. ALERTAS
# =====================================================================
h1("9. Alertas de inventario")
para("El sistema genera alertas automáticas para ayudarle a actuar a tiempo:")
vinetas([
    "Stock mínimo: cuando un producto alcanza o baja de su nivel mínimo.",
    "Por vencer: cuando un producto se acerca a su fecha de vencimiento.",
])
pasos([
    "Abra «Alertas» en el menú lateral (o la campana en la barra superior).",
    "Revise la lista de alertas activas.",
    "Cuando atienda una alerta (por ejemplo, repone el stock), presione «Resolver» para archivarla.",
])
figura("alertas_de_stock_y_vencimiento", "Alertas de stock y vencimiento.")

# =====================================================================
# 10. REPORTES
# =====================================================================
h1("10. Reportes gerenciales")
para("El módulo de Reportes presenta indicadores clave (KPIs) para la toma de decisiones "
     "y permite exportar la información.")
pasos([
    "Abra «Reportes» en el menú lateral.",
    "Consulte los indicadores: total de productos, alertas activas, movimientos del día y "
    "existencias por categoría.",
    "Para respaldar o analizar los datos, use «Exportar inventario» o «Exportar movimientos»; "
    "se descargará un archivo CSV compatible con Excel.",
])
figura("reportes_gerenciales", "Reportes gerenciales e indicadores.")

# =====================================================================
# 11. CONFIGURACIÓN Y RESPALDO
# =====================================================================
h1("11. Configuración y respaldo")
para("En Configuración puede exportar la información como respaldo, consultar los conteos "
     "del sistema y ajustar preferencias.")
pasos([
    "Abra «Configuración» desde el menú de su cuenta (barra superior) o el menú lateral.",
    "Use «Exportar inventario» / «Exportar movimientos» para guardar respaldos en CSV.",
    "Ajuste las preferencias disponibles y presione «Guardar configuración».",
])
figura("respaldo_y_configuraci_n", "Respaldo y configuración del sistema.")
nota("Se recomienda exportar un respaldo de forma periódica (por ejemplo, semanalmente).")

# =====================================================================
# 12. USUARIOS Y ROLES
# =====================================================================
h1("12. Usuarios y roles (solo administradores)")
para("El sistema maneja tres roles, cada uno con su nivel de acceso:")
vinetas([
    "Administrador: acceso completo; además, gestiona los usuarios del sistema.",
    "Gerente: opera el sistema y consulta reportes.",
    "Operario: registra productos y movimientos en el día a día.",
])
para("Para gestionar usuarios (solo administradores):")
pasos([
    "Abra «Usuarios» (o desde Configuración → «Gestionar roles y usuarios»).",
    "Para crear un usuario, complete nombre, correo, contraseña y rol, y guarde.",
    "Para retirar un acceso, use la opción de desactivar al usuario correspondiente.",
])

# =====================================================================
# 13. PREGUNTAS FRECUENTES
# =====================================================================
h1("13. Preguntas frecuentes y solución de problemas")

def faq(preg, resp):
    p = doc.add_paragraph()
    r = p.add_run(preg); r.font.bold = True; r.font.color.rgb = NAVY
    para(resp)

faq("El teléfono dice «sesión expirada» al tomar la foto.",
    "El código QR caduca por seguridad. Recargue la pantalla «Captura IA» en la "
    "computadora para generar un código nuevo y vuelva a escanearlo.")
faq("La cámara del teléfono no se abre.",
    "Permita el acceso a la cámara cuando el navegador lo solicite. Si no aparece la "
    "cámara, use la opción de seleccionar una foto desde la galería.")
faq("El sistema me devolvió al inicio de sesión.",
    "Su sesión expiró por inactividad. Vuelva a iniciar sesión con sus credenciales.")
faq("No puedo registrar una salida.",
    "Verifique que el producto exista en el inventario y que haya existencias suficientes; "
    "el sistema no permite stock negativo.")
faq("El reconocimiento no identificó bien el producto.",
    "Corrija el nombre en la pantalla de confirmación antes de guardar. Tome la foto con "
    "buena iluminación y enfocando la etiqueta del producto.")

doc.save("Manual_de_Usuario.docx")
print("Listo: Manual_de_Usuario.docx")
